import xlrd3
import xlrd
import os
import copy
import openpyxl
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
import shutil

try:
    Path = os.getcwd()
    fileName = Path + "\\" + "ABC.docx"
    print("'{0}'".format(fileName))

    document = Document()
    document.styles['Normal'].font.name = u'宋体'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    document.styles['Normal'].font.size = Pt(18)

    p = document.add_paragraph()
    paragraphHead = p.add_run('开封大学学业预警通知书')
    paragraphHead.bold = True
    paragraphHead.font.size = Pt(16)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = document.add_paragraph()
    paragraphSecondHead = p.add_run('院部（公章）：                                签发日期：')
    paragraphSecondHead.font.size = Pt(12)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #table 1
    table = document.add_table(rows=1, cols=6, style='Table Grid')
    #table.high = Inches(5)
    table.cell(0, 0).width = Inches(1)
    table.cell(0, 1).width = Inches(3)
    table.cell(0, 2).width = Inches(1)
    table.cell(0, 3).width = Inches(2)
    table.cell(0, 4).width = Inches(1)
    table.cell(0, 5).width = Inches(3)

    tabP = table.cell(0, 0).paragraphs[0]
    tableContent = tabP.add_run('班级')
    tableContent.font.size = Pt(12)
    tabP.alignment = WD_ALIGN_PARAGRAPH.CENTER

    tabP = table.cell(0, 1).paragraphs[0]
    tableContent = tabP.add_run('AAAAA')
    tableContent.font.size = Pt(12)
    tabP.alignment = WD_ALIGN_PARAGRAPH.CENTER

    tabP = table.cell(0, 2).paragraphs[0]
    tableContent = tabP.add_run('姓名')
    tableContent.font.size = Pt(12)
    tabP.alignment = WD_ALIGN_PARAGRAPH.CENTER

    tabP = table.cell(0, 3).paragraphs[0]
    tableContent = tabP.add_run('BBBB')
    tableContent.font.size = Pt(12)
    tabP.alignment = WD_ALIGN_PARAGRAPH.CENTER

    tabP = table.cell(0, 4).paragraphs[0]
    tableContent = tabP.add_run('学号')
    tableContent.font.size = Pt(12)
    tabP.alignment = WD_ALIGN_PARAGRAPH.CENTER

    tabP = table.cell(0, 5).paragraphs[0]
    tableContent = tabP.add_run('CCCC')
    tableContent.font.size = Pt(12)
    tabP.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #table 2
    table = document.add_table(rows=1, cols=1,  style='Table Grid')
    #table.high = Inches(5)
    tabP = table.cell(0, 0).paragraphs[0]
    tableContent = tabP.add_run('未通过课程/环节')
    tableContent.font.size = Pt(14)
    tabP.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #table 3
    table = document.add_table(rows=6, cols=8, style='Table Grid')
    table.cell(0, 0).width = Inches(1.5)
    table.cell(0, 1).width = Inches(1.5)
    table.cell(0, 2).width = Inches(1)
    table.cell(0, 3).width = Inches(1)
    table.cell(0, 4).width = Inches(1.5)
    table.cell(0, 5).width = Inches(1.5)
    table.cell(0, 6).width = Inches(1)
    table.cell(0, 7).width = Inches(1)

    tabP = table.cell(0, 0).paragraphs[0]
    tableContent = tabP.add_run('学年学期')
    tableContent.font.size = Pt(12)

    tabP = table.cell(0, 1).paragraphs[0]
    tableContent = tabP.add_run('课程、环节名称')
    tableContent.font.size = Pt(12)

    tabP = table.cell(0, 2).paragraphs[0]
    tableContent = tabP.add_run('成绩')
    tableContent.font.size = Pt(12)
    tabP.alignment = WD_ALIGN_PARAGRAPH.CENTER

    tabP = table.cell(0, 3).paragraphs[0]
    tableContent = tabP.add_run('学分')
    tableContent.font.size = Pt(12)

    tabP = table.cell(0, 4).paragraphs[0]
    tableContent = tabP.add_run('学年学期')
    tableContent.font.size = Pt(12)

    tabP = table.cell(0, 5).paragraphs[0]
    tableContent = tabP.add_run('课程、环节名称')
    tableContent.font.size = Pt(12)

    tabP = table.cell(0, 6).paragraphs[0]
    tableContent = tabP.add_run('成绩')
    tableContent.font.size = Pt(12)
    tabP.alignment = WD_ALIGN_PARAGRAPH.CENTER

    tabP = table.cell(0, 7).paragraphs[0]
    tableContent = tabP.add_run('学分')
    tableContent.font.size = Pt(12)

    #table 4
    table = document.add_table(rows=4, cols=1, style='Table Grid')
    tabP = table.cell(0, 0).paragraphs[0]
    tableContent = tabP.add_run('告知意见')
    tableContent.font.size = Pt(14)
    tabP.alignment = WD_ALIGN_PARAGRAPH.CENTER

    tabP1 = table.cell(1, 0).paragraphs[0]
    tableContent = tabP1.add_run('该同学因不及格学分达到学业预警标准，特告知学生及家长，以作警示，如果学生以后再出现不及格现象，可能造成以下后果：')
    tableContent.font.size = Pt(12)
    tabP1.paragraph_format.first_line_indent = 305000

    tabP2 = table.cell(1, 0).add_paragraph()
    tableContent = tabP2.add_run('一、学生在校期间，累计不及格课程学分超过25学分，应予退学。')
    tableContent.font.size = Pt(12)
    tabP2.paragraph_format.first_line_indent = 305000

    tabP3 = table.cell(1, 0).add_paragraph()
    tableContent = tabP3.add_run('二、学生在学校基本学制年限内，未达到毕业要求的，按结业处理。')
    tableContent.font.size = Pt(12)
    tabP3.paragraph_format.first_line_indent = 305000

    tabP4 = table.cell(1, 0).add_paragraph()
    tableContent = tabP4.add_run('学生签名:')
    tableContent.font.size = Pt(12)
    tabP4.paragraph_format.first_line_indent = 4727500

    tabP = table.cell(2, 0).paragraphs[0]
    tableContent = tabP.add_run('家长意见：\n')
    tableContent.font.size = Pt(12)
    tableContent = tabP.add_run('1、学生目前学业情况及将来可能造成后果是否已知悉：______________________\n')
    tableContent.font.size = Pt(12)
    tableContent = tabP.add_run('2、家长是否已认真阅读《开封大学学籍管理规定》_______________________\n')
    tableContent.font.size = Pt(12)
    tableContent = tabP.add_run('3、家长意见及建议（可另附纸）：\n\n\n')
    tableContent.font.size = Pt(12)
    tableContent = tabP.add_run('家长签名：\n')
    tableContent.font.size = Pt(12)
    tableContent = tabP.add_run('家长通信地址及联系方式：\n')
    tableContent.font.size = Pt(12)

    tabP = table.cell(3, 0).paragraphs[0]
    tableContent = tabP.add_run('院部告知情况记录：\n\n\n\n')
    tableContent.font.size = Pt(12)
    tableContent = tabP.add_run('负责人签名：')
    tableContent.font.size = Pt(12)

    p = document.add_paragraph()
    paragraphSecondHead = p.add_run('本表一式两份，一份学生留存、一份院部留存')
    paragraphSecondHead.font.size = Pt(10.5)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.save(fileName)


except BaseException as err:
    print("{0}".format(err))



