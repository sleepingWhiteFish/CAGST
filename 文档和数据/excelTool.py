# -!- coding: utf-8 -!-
"""
Collect multiple execels information, summarize to one execel.
version:v1.0
plateform and dependent version information:
    python 3.10.0
    xlrd3 1.1.0
    openpyxl 3.0.9
    python-docx 0.8.11
Created by LC on 2021.10.9
"""

"""
默认情况读取脚本同级目录下的".xlsx"文件。
直接输入文件名（不带后缀）时，默认搜索脚本所在目录。
可以带路径+文件名，文件名无需加后缀，比如：”D:\AA\BB\CC.xlsx“，需要直接输入“D:\AA\BB\CC“。
生成名为"容灾配置"的目录,生成在输入文件名同级目录中。
"""

import xlrd3
import xlrd
import os
import copy
import openpyxl
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
import shutil

class notPassPeopleClassExlInfo(object):
    className = ""
    notPassInfoStartRow = -1
    studentIdColNum = -1
    studentNameColNum = -1
    creditColNum = -1
    courseTypeColNum = -1
    courseColNum = -1
    assessmentMethodColNum = -1
    achievementColNum = -1
    semesterColNum = -1
    specialColNum = -1

    def _init_(self):
        self.className = ""
        self.notPassInfoStartRow = -1
        self.studentIdColNum = -1
        self.studentNameColNum = -1
        self.creditColNum = -1
        self.courseTypeColNum = -1
        self.courseColNum = -1
        self.assessmentMethodColNum = -1
        self.achievementColNum = -1
        self.semesterColNum = -1
        self.specialColNum = -1

    def cleanData(self):
        self.className = ""
        self.notPassInfoStartRow = -1
        self.studentIdColNum = -1
        self.studentNameColNum = -1
        self.creditColNum = -1
        self.courseTypeColNum = -1
        self.courseColNum = -1
        self.assessmentMethodColNum = -1
        self.achievementColNum = -1
        self.semesterColNum = -1
        self.specialColNum = -1

    def setClassName(self, inClassName):
        self.className = inClassName

    def getClassName(self):
        return self.className

    def setStartRow(self, inStartRow):
        self.notPassInfoStartRow = inStartRow

    def getStartRow(self):
        return self.notPassInfoStartRow

    def setStudentIdColNum(self, inStudentIdColNum):
        self.studentIdColNum = inStudentIdColNum

    def getStudentIdColNum(self):
        return self.studentIdColNum

    def setStudentNameColNum(self, inStudentNameColNum):
        self.studentNameColNum = inStudentNameColNum

    def getStudentNameColNum(self):
        return self.studentNameColNum

    def setCreditColNum(self, inCreditColNum):
        self.creditColNum = inCreditColNum

    def getCreditColNum(self):
        return self.creditColNum

    def setCourseTypeColNum(self, inCourseTypeColNum):
        self.courseTypeColNum = inCourseTypeColNum

    def getCourseTypeColNum(self):
        return self.courseTypeColNum

    def setCourseColNum(self, inCourseColNum):
        self.courseColNum = inCourseColNum

    def getCourseColNum(self):
        return self.courseColNum

    def setAssessmentMethodColNum(self, inAssessmentMethodColNum):
        self.assessmentMethodColNum = inAssessmentMethodColNum

    def getAssessmentMethodColNum(self):
        return self.assessmentMethodColNum

    def setAchievementColNum(self, inAchievementColNum):
        self.achievementColNum = inAchievementColNum

    def getAchievementColNum(self):
        return self.achievementColNum

    def setSemesterColNum(self, inSemesterColNum):
        self.semesterColNum = inSemesterColNum

    def getSemesterColNum(self):
        return self.semesterColNum

    def setSpecialColNum(self, inSpecialColNum):
        self.specialColNum = inSpecialColNum

    def getSpecialColNum(self):
        return self.specialColNum

class remakeRecordDoc(object):
    semester = ""
    className = ""
    studentId = 0
    studentName = ""
    courseName = ""
    courseId = 0
    achievement = 0
    credit = 0.0
    special = ""
    haveDataFlag = 0

    def _init_(self):
        self.semester = ""
        self.className = ""
        self.studentId = 0
        self.studentName = ""
        self.courseName = ""
        self.courseId = 0
        self.achievement = 0
        self.credit = 0.0
        self.special = ""
        self.haveDataFlag = 0

    def cleanData(self):
        self.semester = ""
        self.className = ""
        self.studentId = 0
        self.studentName = ""
        self.courseName = ""
        self.courseId = 0
        self.achievement = 0
        self.credit = 0.0
        self.special = ""
        self.haveDataFlag = 0

    def setsemester(self, insemester):
        self.semester = insemester

    def getsemester(self):
        return self.semester

    def setclassName(self, inclassName):
            self.className = inclassName

    def getclassName(self):
            return self.className

    def setstudentId(self, instudentId):
            self.studentId = instudentId

    def getstudentId(self):
            return self.studentId

    def setstudentName(self, instudentName):
            self.studentName = instudentName

    def getstudentName(self):
            return self.studentName

    def setcourseName(self, incourseName):
            self.courseName = incourseName

    def getcourseName(self):
            return self.courseName

    def setcourseId(self, incourseId):
            self.courseId = incourseId

    def getcourseId(self):
            return self.courseId

    def setachievement(self, inachievement):
            self.achievement = inachievement

    def getachievement(self):
            return self.achievement

    def setcredit(self, incredit):
            self.credit = incredit

    def getcredit(self):
            return self.credit

    def setspecial(self, inspecial):
            self.special = inspecial

    def getspecial(self):
            return self.special

    def sethaveDataFlag(self, inhaveDataFlag):
            self.haveDataFlag += inhaveDataFlag

    def gethaveDataFlag(self):
            return self.haveDataFlag

def generateNotification(notifyInfo, targetPath):
    try:
        targetFilePath = targetPath + "重修通知单"
        fileName = notifyInfo.getclassName() + "班" + notifyInfo.getstudentName() + notifyInfo.getcourseName()
        if not os.path.exists(targetFilePath):
            os.makedirs(targetFilePath)

        fileName = targetFilePath + "\\" + fileName + ".docx"

        document = Document()
        document.styles['Normal'].font.name = u'宋体'
        document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        document.styles['Normal'].font.size = Pt(18)

        p = document.add_paragraph()
        paragraphHead = p.add_run('开封大学学生课程重修登记表')
        paragraphHead.bold = True
        paragraphHead.font.size = Pt(18)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p = document.add_paragraph()
        paragraphSecondHead = p.add_run('学生所在院（院）： 财政经济学院     2019-2020  学年 第二学期')
        paragraphSecondHead.font.size = Pt(14)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # table 1
        table = document.add_table(rows=2, cols=4, style='Table Grid')
        # table.high = Inches(5)
        tabP = table.cell(0, 0).paragraphs[0]
        tableContent = tabP.add_run('班  级')
        tableContent.font.size = Pt(12)
        tabP.alignment = WD_ALIGN_PARAGRAPH.CENTER

        tabP = table.cell(0, 1).paragraphs[0]
        tableContent = tabP.add_run('学  号')
        tableContent.font.size = Pt(12)
        tabP.alignment = WD_ALIGN_PARAGRAPH.CENTER

        tabP = table.cell(0, 2).paragraphs[0]
        tableContent = tabP.add_run('姓  名')
        tableContent.font.size = Pt(12)
        tabP.alignment = WD_ALIGN_PARAGRAPH.CENTER

        tabP = table.cell(0, 3).paragraphs[0]
        tableContent = tabP.add_run('联系方式')
        tableContent.font.size = Pt(12)
        tabP.alignment = WD_ALIGN_PARAGRAPH.CENTER

        tabP = table.cell(1, 0).paragraphs[0]
        tableContent = tabP.add_run(notifyInfo.getclassName() + "班")
        tableContent.font.size = Pt(12)
        tabP.alignment = WD_ALIGN_PARAGRAPH.CENTER

        tabP = table.cell(1, 1).paragraphs[0]
        tableContent = tabP.add_run(notifyInfo.getstudentId())
        tableContent.font.size = Pt(12)
        tabP.alignment = WD_ALIGN_PARAGRAPH.CENTER

        tabP = table.cell(1, 2).paragraphs[0]
        tableContent = tabP.add_run(notifyInfo.getstudentName())
        tableContent.font.size = Pt(12)
        tabP.alignment = WD_ALIGN_PARAGRAPH.CENTER

        tabP = table.cell(1, 3).paragraphs[0]
        tableContent = tabP.add_run('')
        tableContent.font.size = Pt(12)
        tabP.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # table 2
        table = document.add_table(rows=2, cols=6, style='Table Grid')
        # table.high = Inches(5)
        tabP = table.cell(0, 0).paragraphs[0]
        tableContent = tabP.add_run('课程承担单位')
        tableContent.font.size = Pt(12)
        tabP.alignment = WD_ALIGN_PARAGRAPH.CENTER

        tabP = table.cell(0, 1).paragraphs[0]
        tableContent = tabP.add_run('课程代码')
        tableContent.font.size = Pt(12)
        tabP.alignment = WD_ALIGN_PARAGRAPH.CENTER

        tabP = table.cell(0, 2).paragraphs[0]
        tableContent = tabP.add_run('课程名称')
        tableContent.font.size = Pt(12)
        tabP.alignment = WD_ALIGN_PARAGRAPH.CENTER

        tabP = table.cell(0, 3).paragraphs[0]
        tableContent = tabP.add_run('初修成绩')
        tableContent.font.size = Pt(12)
        tabP.alignment = WD_ALIGN_PARAGRAPH.CENTER

        tabP = table.cell(0, 4).paragraphs[0]
        tableContent = tabP.add_run('学时/学分')
        tableContent.font.size = Pt(12)
        tabP.alignment = WD_ALIGN_PARAGRAPH.CENTER

        tabP = table.cell(0, 5).paragraphs[0]
        tableContent = tabP.add_run('备注')
        tableContent.font.size = Pt(12)
        tabP.alignment = WD_ALIGN_PARAGRAPH.CENTER

        tabP = table.cell(1, 0).paragraphs[0]
        tableContent = tabP.add_run('')
        tableContent.font.size = Pt(12)
        tabP.alignment = WD_ALIGN_PARAGRAPH.CENTER

        tabP = table.cell(1, 1).paragraphs[0]
        tableContent = tabP.add_run(notifyInfo.getcourseId())
        tableContent.font.size = Pt(12)
        tabP.alignment = WD_ALIGN_PARAGRAPH.CENTER

        tabP = table.cell(1, 2).paragraphs[0]
        tableContent = tabP.add_run(notifyInfo.getcourseName())
        tableContent.font.size = Pt(12)
        tabP.alignment = WD_ALIGN_PARAGRAPH.CENTER

        tabP = table.cell(1, 3).paragraphs[0]
        tableContent = tabP.add_run(str(notifyInfo.getachievement()))
        tableContent.font.size = Pt(12)
        tabP.alignment = WD_ALIGN_PARAGRAPH.CENTER

        tabP = table.cell(1, 4).paragraphs[0]
        value = str(7.5 * notifyInfo.getcredit()) + '/' + str(notifyInfo.getcredit())
        tableContent = tabP.add_run(value)
        tableContent.font.size = Pt(12)
        tabP.alignment = WD_ALIGN_PARAGRAPH.CENTER

        tabP = table.cell(1, 5).paragraphs[0]
        tableContent = tabP.add_run(notifyInfo.getspecial())
        tableContent.font.size = Pt(12)
        tabP.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # table 3
        table = document.add_table(rows=4, cols=1, style='Table Grid')
        tabP = table.cell(0, 0).paragraphs[0]
        tableContent = tabP.add_run('院部意见：\n\n\n\n\n')
        tableContent.font.size = Pt(12)

        tabP = table.cell(1, 0).paragraphs[0]
        tableContent = tabP.add_run('承担单位意见：\n\n\n\n\n')
        tableContent.font.size = Pt(12)

        tabP = table.cell(2, 0).paragraphs[0]
        tableContent = tabP.add_run('考勤／辅导记录：')
        tableContent.font.size = Pt(12)
        tabP.alignment = WD_ALIGN_PARAGRAPH.CENTER

        tabP = table.cell(3, 0).paragraphs[0]
        tableContent = tabP.add_run('\n\n\n\n\n')
        tableContent.font.size = Pt(12)

        # table 4
        table = document.add_table(rows=3, cols=2, style='Table Grid')
        # table width all 10 = 3 + 7 ,everyone need to set
        # table.cell(0, 0).width = Inches(1.5)
        # table.cell(0, 1).width = Inches(8.5)
        tabP = table.cell(0, 0).paragraphs[0]
        tableContent = tabP.add_run('考试时间')
        tableContent.font.size = Pt(12)
        tabP.alignment = WD_ALIGN_PARAGRAPH.CENTER

        tabP = table.cell(0, 1).paragraphs[0]
        tableContent = tabP.add_run('总评成绩')
        tableContent.font.size = Pt(12)
        tabP.alignment = WD_ALIGN_PARAGRAPH.CENTER

        tabP = table.cell(1, 1).paragraphs[0]
        tableContent = tabP.add_run('总评成绩______=平时成绩_____×_____% +末考成绩_____×_____%')
        tableContent.font.size = Pt(12)
        tabP.alignment = WD_ALIGN_PARAGRAPH.CENTER

        tabP = table.cell(2, 0).paragraphs[0]
        tableContent = tabP.add_run('重修形式')
        tableContent.font.size = Pt(12)
        tabP.alignment = WD_ALIGN_PARAGRAPH.CENTER

        tabP = table.cell(2, 1).paragraphs[0]
        tableContent = tabP.add_run('自学辅导/插班重修/单独编班')
        tableContent.font.size = Pt(12)
        tabP.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # table 5
        table = document.add_table(rows=1, cols=1, style='Table Grid')
        tabP = table.cell(0, 0).paragraphs[0]
        tableContent = tabP.add_run('任课教师签字：\n\n')
        tableContent.font.size = Pt(12)

        p = document.add_paragraph()
        paragraphSecondHead = p.add_run('备注：本表一式两份，一份任课教师填写，末考后交院部，一份学生所在院部留存备查。')
        paragraphSecondHead.font.size = Pt(10.5)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        document.save(fileName)
        return True
    except BaseException as err:
        print("{0}".format(err))
        return False

def retakeDealExl(firstFilePathAndName='系统导出未通过', secondFilePathAndName='教师任务表一', thirdFilePathAndName='教师任务表二'):
    """Reade xlsx file and get date"""
    try:
        # first file and path
        if ("\\" in firstFilePathAndName):
            charIndex = firstFilePathAndName.rfind("\\")
            charIndex += 1
            path = firstFilePathAndName[0:charIndex]
            firstFileName = firstFilePathAndName + '.xlsx'
            if not os.path.exists(firstFileName):
                firstFileName = firstFilePathAndName + '.xls'

            if not os.path.exists(firstFileName):
                print("错误：找不到文件--{0}".format(firstFileName))
                return
        else:
            path = os.getcwd() + "\\"
            firstFileName = path + firstFilePathAndName + '.xlsx'
            if not os.path.exists(firstFileName):
                firstFileName = firstFilePathAndName + '.xls'

            if not os.path.exists(firstFileName):
                print("错误：找不到文件--{0}".format(firstFileName))
                return

        resultDir = path + "重修通知单"
        if os.path.exists(resultDir):
            shutil.rmtree(resultDir)

        if ("xlsx" in firstFileName):
            notPassBook = xlrd3.open_workbook(firstFileName)
        elif ("xls" in firstFileName):
            notPassBook = xlrd.open_workbook(firstFileName)
        else:
            print("错误：文件不是excel格式--{0}".format(firstFileName))
            return

        notPassSheetsNum = notPassBook.nsheets
        notPassSheetIndex = 0
        classExl = notPassPeopleClassExlInfo()
        oneRemakeNotify = remakeRecordDoc()

        # sheet loop
        while (notPassSheetIndex < notPassSheetsNum):
            currentSheet = notPassBook.sheet_by_index(notPassSheetIndex)
            rowNum = currentSheet.nrows
            colNum = currentSheet.ncols
            rowIndex = 0
            while (rowIndex < rowNum):
                colIndex = 0
                oneRemakeNotify.cleanData()
                while (colIndex < colNum):
                    currentColValue = currentSheet.cell_value(rowx=rowIndex, colx=colIndex)
                    if ("行政班级" in str(currentColValue)):
                        charIndex = currentColValue.find("行政班级：")
                        classNameStart = charIndex + len("行政班级：")
                        classNameLen = currentColValue[classNameStart:].find(" ")
                        currentClassName = currentColValue[classNameStart:classNameStart + classNameLen]
                        classExl.cleanData()
                        classExl.setClassName(currentClassName)
                        classExl.setStartRow(rowIndex + 1)
                        break
                    else:
                        if (classExl.getClassName() != ""):
                            if(classExl.getStartRow() == rowIndex):
                                if(currentColValue == '学号'):
                                    classExl.setStudentIdColNum(colIndex)
                                elif(currentColValue == '姓名'):
                                    classExl.setStudentNameColNum(colIndex)
                                elif (currentColValue == '学分'):
                                    classExl.setCreditColNum(colIndex)
                                elif (currentColValue == '课程类别'):
                                    classExl.setCourseTypeColNum(colIndex)
                                elif (currentColValue == '课程/环节'):
                                    classExl.setCourseColNum(colIndex)
                                elif (currentColValue == '考核方式'):
                                    classExl.setAssessmentMethodColNum(colIndex)
                                elif (currentColValue == '成绩'):
                                    classExl.setAchievementColNum(colIndex)
                                elif (currentColValue == '学年学期'):
                                    classExl.setSemesterColNum(colIndex)
                                elif ('特殊' in currentColValue):
                                    classExl.setSpecialColNum(colIndex)

                                if (colIndex == colNum - 1):
                                    if(classExl.getStudentIdColNum() == -1 or
                                    classExl.getStudentNameColNum() == -1 or
                                    classExl.getCreditColNum() == -1 or
                                    classExl.getCourseTypeColNum() == -1 or
                                    classExl.getCourseColNum() == -1 or
                                    classExl.getAssessmentMethodColNum() == -1 or
                                    classExl.getAchievementColNum() == -1 or
                                    classExl.getSemesterColNum() == -1 or
                                    classExl.getSpecialColNum() == -1):
                                        print("错误：班级--'{0}'缺少列名\n".format(classExl.getClassName()))
                                        return
                            else:
                                oneRemakeNotify.setclassName(classExl.getClassName())
                                if (colIndex == classExl.getStudentIdColNum()):
                                    if ("" != currentColValue):
                                        oneRemakeNotify.setstudentId(currentColValue)
                                        oneRemakeNotify.sethaveDataFlag(1)
                                    else:
                                        print("错误：行号--'{0}'，缺少学号数据\n".format(rowIndex))
                                        return
                                elif (colIndex == classExl.getStudentNameColNum()):
                                    if ("" != currentColValue):
                                        oneRemakeNotify.setstudentName(currentColValue)
                                        oneRemakeNotify.sethaveDataFlag(2)
                                    else:
                                        print("错误：行号--'{0}'，缺少学生姓名数据\n".format(rowIndex))
                                        return
                                elif (colIndex == classExl.getCreditColNum()):
                                    if ("" != currentColValue):
                                        oneRemakeNotify.setcredit(currentColValue)
                                        oneRemakeNotify.sethaveDataFlag(4)
                                    else:
                                        print("错误：行号--'{0}'，缺少学分数据\n".format(rowIndex))
                                        return
                                elif (colIndex == classExl.getCourseColNum()):
                                    if ("" != currentColValue):
                                        couseNameStart = currentColValue.find('[')
                                        couseNameEnd = currentColValue.find(']')
                                        courseId = currentColValue[couseNameStart + 1:couseNameEnd]
                                        courseName = currentColValue[couseNameEnd + 1:]
                                        oneRemakeNotify.setcourseName(courseName)
                                        oneRemakeNotify.sethaveDataFlag(8)
                                        oneRemakeNotify.setcourseId(courseId)
                                        oneRemakeNotify.sethaveDataFlag(16)
                                    else:
                                        print("错误：行号--'{0}'，缺少课程名称代码数据\n".format(rowIndex))
                                        return
                                elif (colIndex == classExl.getAchievementColNum()):
                                    if ("" != currentColValue):
                                        oneRemakeNotify.setachievement(currentColValue)
                                elif (colIndex == classExl.getSpecialColNum()):
                                    if ("" != currentColValue):
                                        oneRemakeNotify.setspecial(currentColValue)
                                elif (colIndex == classExl.getSemesterColNum()):
                                    if ("" != currentColValue):
                                        oneRemakeNotify.setsemester(currentColValue)
                                        oneRemakeNotify.sethaveDataFlag(32)
                                    else:
                                        print("错误：行号--'{0}'，缺少学年学期数据\n".format(rowIndex))
                                        return
                    colIndex += 1
                if(63 == oneRemakeNotify.gethaveDataFlag()):
                    print("第{0}行--姓名：{1}--课程：{2}\n".format(rowIndex + 1, oneRemakeNotify.getstudentName(), oneRemakeNotify.getcourseName()))
                    if(False == generateNotification(oneRemakeNotify, path)):
                        print("第'{0}行数据生成通知失败\n".format(rowIndex + 1))
                rowIndex += 1
        return
    except BaseException as err:
        print("{0}".format(err))
        return

def retakeMainFunc():
    try:
        while True:
            retakeInpute = input("输入 'd' or 'D' 用默认文件名。默认名为：系统导出未通过、教师任务表一、教师任务表二\n\
输入 's' or 'S' 进入重修输入文件名录入。\n\
输入 'q' or 'Q' 退到上一层\n")
            # noinspection PyRedundantParentheses
            if (retakeInpute == "d" or "D" == retakeInpute):
                retakeDealExl()
            elif (retakeInpute == "s" or retakeInpute == "S"):
                retakeFirstExelName = input("输入重修系统导出未通过表格路径加文件名，不带后缀。\n")
                retakeSecondExelName = input("输入教师任务表一路径加文件名，不带后缀。\n")
                retakeThirdExelName = input("输入教师任务表二路径加文件名，不带后缀。\n")
                retakeDealExl(retakeFirstExelName, retakeSecondExelName, retakeThirdExelName)
            elif (retakeInpute == "q" or retakeInpute == "Q"):
                return
            else:
                print("输入未知，重新输入。\n")
    except BaseException as err:
        print("{0}".format(err))
        return

try:
    while True:
        userInpute = input("\n输入 'Q' or 'q' 退出。\n\
输入 'cx' 进入重修名单整理功能。\n")

        if (userInpute == "q" or userInpute == "Q"):
            break
        elif (userInpute == "cx"):
            print("\n进入重修整理功能......\n")
            retakeMainFunc()
except  BaseException as err:
    print("{0}".format(err))
